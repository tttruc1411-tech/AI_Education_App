from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = Document()
    
    # Title Section
    title = doc.add_heading('AI Coding Lab: Architecture & Implementation Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Design guide for the Function Library and Learning Hub').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Last Updated: April 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Section: High Level Architecture ---
    doc.add_heading('1. High-Level Architecture', level=1)
    doc.add_paragraph(
        "The Function Library is a data-driven system designed to abstract complex computer vision and robotics logic into "
        "simple, draggable blocks for students. It separates purely UI presentation from technical execution and metadata."
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Location'
    hdr_cells[2].text = 'Purpose'
    
    layers = [
        ('Frontend (UI)', 'src/modules/function_library.py', 'Renders the sidebar, category headers, and draggable cards.'),
        ('Registry (Data)', 'src/modules/library/definitions.py', 'Contains the master LIBRARY_FUNCTIONS dictionary defining every block.'),
        ('Backend (Impl)', 'src/modules/library/functions/*.py', 'Actual Python implementations for AI, Vision, and Robotics.')
    ]
    
    for layer, loc, purpose in layers:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = loc
        row_cells[2].text = purpose

    # --- Section: How Blocks Work? ---
    doc.add_heading('2. How a Block Works (Logic & Metadata)', level=1)
    doc.add_paragraph(
        "Every function in the library is described by a specific JSON-like dictionary entry in definitions.py. "
        "This registry is what 'teaches' the editor how to handle code injection and documentation."
    )
    
    doc.add_heading('2.1 Key Metadata Fields', level=2)
    fields = [
        ('import_statement', 'The line added to the top of the editor (e.g., from src.modules... import foo)'),
        ('usage', 'The snippet injected at the cursor position (e.g., res = foo(x))'),
        ('params', 'Describes arguments for the sidebar documentation panel.'),
        ('source_func', 'Points to the exact function name in the source file for live code preview.'),
        ('source_module', 'Points to the .py module containing the function implementation.')
    ]
    
    for field, desc in fields:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(field)
        run.bold = True
        p.add_run(f": {desc}")

    # --- Section: Implementation Guide ---
    doc.add_heading('3. Implementation Guide: Adding a New Function', level=1)
    doc.add_paragraph("Step-by-step instructions to expand the library.")
    
    # Step 1
    doc.add_heading('Step 1: Write the Backend Logic', level=2)
    doc.add_paragraph(
        "Place your actual Python code in a file under src/modules/library/functions/. "
        "For example, create sensors.py if you are adding environmental sensors."
    )
    code_block = doc.add_paragraph(style='Normal')
    run = code_block.add_run("def Read_Moisture(pin):\n    # Internal logic here\n    return 45.0")
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)

    # Step 2
    doc.add_heading('Step 2: Register in definitions.py', level=2)
    doc.add_paragraph("Open definitions.py and find the target Category (or create a new one). Add the metadata entry:")
    code_block = doc.add_paragraph(style='Normal')
    run = code_block.add_run(
        '\"Read_Moisture\": {\n'
        '    \"desc\": \"Get soil moisture level\",\n'
        '    \"usage\": \"val = Read_Moisture(1)\",\n'
        '    \"import_statement\": \"from src.modules.library.functions.sensors import Read_Moisture\",\n'
        '    \"source_func\": \"Read_Moisture\",\n'
        '    \"source_module\": \"src.modules.library.functions.sensors\"\n'
        '}'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # --- Section: New Group ---
    doc.add_heading('4. Implementation Guide: Adding a New Group (Category)', level=1)
    doc.add_paragraph(
        "To add a new vertical group (like the one you created for 'Robotics'):"
    )
    doc.add_paragraph(
        "1. Create the backend script script (e.g. robotics.py).\n"
        "2. In definitions.py, insert a new top-level key into the LIBRARY_FUNCTIONS dictionary.\n"
        "3. Provide a 'color' (Hex code) and an 'icon' (Emoji or Path).\n"
        "4. The order of keys in LIBRARY_FUNCTIONS determines the vertical order in the UI.",
        style='List Number'
    )

    # Final Notes
    doc.add_heading('5. UI Component Hierarchy', level=1)
    doc.add_paragraph("CategoryHeader -> Container (Grouped Blocks) -> DraggableFunctionBlock -> FunctionInfoPanel (Collapsible Info).")

    # Save
    doc_path = os.path.join(os.getcwd(), 'Design_and_Architecture_Function_Library.docx')
    doc.save(doc_path)
    print(f"Document saved to: {doc_path}")

if __name__ == "__main__":
    create_doc()
